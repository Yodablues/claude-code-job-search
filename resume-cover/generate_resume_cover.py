"""
Generate tailored resume and cover letter .docx files from JSON input.

Usage:
    python generate_resume_cover.py --config /path/to/config.json

Config (config.json):
    {
        "name": "Your Name",
        "location": "City, State",
        "email": "you@example.com",
        "phone": "555-555-5555",
        "linkedin": "linkedin.com/in/your-profile",
        "website": "yoursite.com",
        "github": "github.com/you",
        "source_resume": "/path/to/Your_Resume.docx",
        "output_dir": "/path/to/output/directory"
    }

Reads (from output_dir):
    - resume_tailored.json
    - cover_letter_tailored.json

Writes (to output_dir):
    - {Name}_Resume_{company}_{title}.docx
    - {Name}_CoverLetter_{company}_{title}.docx

JSON Schemas
============

resume_tailored.json:
{
    "company": "Acme Corp",
    "short_title": "Staff_FE",
    "summary": "Full text of professional summary...",
    "jobs": [
        {
            "title": "Lead Software Engineer",
            "company_line": "Company Name  |  Mar 2023 - Present  |  City, ST (Remote)",
            "bullets": [
                "Bullet point text...",
                "Another bullet..."
            ]
        }
    ],
    "education": [
        "Bachelor of Science, Computer Science -- University Name  |  2010 - 2014"
    ],
    "skills": [
        {"label": "Frontend", "value": "JavaScript, TypeScript, React..."},
        {"label": "Backend", "value": "Python, Node.js..."},
        {"label": "Tools", "value": "Git, Docker..."}
    ]
}

cover_letter_tailored.json:
{
    "company": "Acme Corp",
    "short_title": "Staff_FE",
    "recipient": "Dear Acme Corp Hiring Team,",
    "opening": "Opening paragraph text...",
    "body_sections": [
        {
            "bold_lead": "Frontend platform leadership at scale. ",
            "text": "Rest of the paragraph..."
        }
    ],
    "closing": "Closing paragraph text...",
    "signoff": "Thank you for your consideration."
}
"""

import argparse
import json
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def load_config(config_path):
    with open(config_path, encoding="utf-8") as f:
        cfg = json.load(f)
    required = ["name", "email", "output_dir"]
    for key in required:
        if key not in cfg:
            print(f"ERROR: config.json missing required key: {key}", file=sys.stderr)
            sys.exit(1)
    return cfg


def set_run(run, size=11, bold=False, color=None):
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_body(doc, text, size=11, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT,
             space_after=2, space_before=0):
    p = doc.add_paragraph()
    p.alignment = align
    p.space_after = Pt(space_after)
    p.space_before = Pt(space_before)
    run = p.add_run(text)
    set_run(run, size=size, bold=bold)
    return p


def add_bullet(doc, text, size=10):
    p = doc.add_paragraph(style="List Bullet")
    p.space_after = Pt(1)
    p.space_before = Pt(1)
    p.clear()
    run = p.add_run(text)
    set_run(run, size=size)


def add_separator(doc):
    p = doc.add_paragraph()
    p.space_after = Pt(0)
    p.space_before = Pt(0)
    run = p.add_run("_" * 80)
    set_run(run, size=8, color=(180, 180, 180))


# -- Resume ------------------------------------------------------------------


def build_resume(data, cfg):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(2)

    for sec in doc.sections:
        sec.top_margin = Inches(0.6)
        sec.bottom_margin = Inches(0.5)
        sec.left_margin = Inches(0.7)
        sec.right_margin = Inches(0.7)

    # Header — name
    add_body(doc, cfg["name"], size=20, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=1)

    # Header — contact line(s)
    contact_parts = []
    if cfg.get("location"):
        contact_parts.append(cfg["location"])
    if cfg.get("linkedin"):
        contact_parts.append(cfg["linkedin"])
    if cfg.get("website"):
        contact_parts.append(cfg["website"])
    if contact_parts:
        add_body(doc, "  |  ".join(contact_parts),
                 size=10, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)

    contact2 = []
    if cfg.get("email"):
        contact2.append(cfg["email"])
    if cfg.get("phone"):
        contact2.append(cfg["phone"])
    if contact2:
        add_body(doc, "  |  ".join(contact2),
                 size=10, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)

    add_separator(doc)

    # Summary
    add_body(doc, "PROFESSIONAL SUMMARY", size=12, bold=True,
             space_before=8, space_after=2)
    add_body(doc, data["summary"], size=10, space_after=4)

    add_separator(doc)

    # Experience
    add_body(doc, "PROFESSIONAL EXPERIENCE", size=12, bold=True,
             space_before=8, space_after=2)

    for job in data["jobs"]:
        p = doc.add_paragraph()
        p.space_before = Pt(6)
        p.space_after = Pt(0)
        r = p.add_run(job["title"])
        set_run(r, size=11, bold=True)

        p2 = doc.add_paragraph()
        p2.space_before = Pt(0)
        p2.space_after = Pt(2)
        r2 = p2.add_run(job["company_line"])
        set_run(r2, size=10, color=(100, 100, 100))

        for b in job["bullets"]:
            add_bullet(doc, b, size=10)

    add_separator(doc)

    # Education
    add_body(doc, "EDUCATION", size=12, bold=True,
             space_before=8, space_after=2)
    for i, edu in enumerate(data["education"]):
        sa = 1 if i < len(data["education"]) - 1 else 4
        add_body(doc, edu, size=10, space_after=sa)

    add_separator(doc)

    # Skills
    add_body(doc, "TECHNICAL SKILLS", size=12, bold=True,
             space_before=8, space_after=2)
    for skill in data["skills"]:
        p = doc.add_paragraph()
        p.space_after = Pt(1)
        p.space_before = Pt(1)
        r1 = p.add_run(f"{skill['label']}: ")
        set_run(r1, size=10, bold=True)
        r2 = p.add_run(skill["value"])
        set_run(r2, size=10)

    return doc


# -- Cover Letter ------------------------------------------------------------


def build_cover_letter(data, cfg):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)

    for sec in doc.sections:
        sec.top_margin = Inches(1)
        sec.bottom_margin = Inches(1)
        sec.left_margin = Inches(1)
        sec.right_margin = Inches(1)

    # Sender info
    add_body(doc, cfg["name"], size=14, bold=True, space_after=0)
    if cfg.get("location"):
        add_body(doc, cfg["location"], size=11, space_after=0)

    contact = []
    if cfg.get("email"):
        contact.append(cfg["email"])
    if cfg.get("phone"):
        contact.append(cfg["phone"])
    if contact:
        add_body(doc, "  |  ".join(contact), size=11, space_after=0)

    links = []
    if cfg.get("website"):
        links.append(cfg["website"])
    if cfg.get("linkedin"):
        links.append(cfg["linkedin"])
    if cfg.get("github"):
        links.append(cfg["github"])
    if links:
        add_body(doc, "  |  ".join(links), size=11, space_after=12)

    # Greeting
    add_body(doc, data["recipient"], size=11, space_after=8, space_before=4)

    # Opening paragraph
    add_body(doc, data["opening"], size=11, space_after=8)

    # Body sections with bold leads
    for section in data["body_sections"]:
        p = doc.add_paragraph()
        p.space_after = Pt(8)
        r1 = p.add_run(section["bold_lead"])
        set_run(r1, size=11, bold=True)
        r2 = p.add_run(section["text"])
        set_run(r2, size=11)

    # Closing
    add_body(doc, data["closing"], size=11, space_after=8)

    # Sign-off
    add_body(doc, data["signoff"], size=11, space_after=12)
    add_body(doc, "Sincerely,", size=11, space_after=0)
    add_body(doc, cfg["name"], size=11, bold=True)

    return doc


# -- Main --------------------------------------------------------------------


def main():
    parser = argparse.ArgumentParser(
        description="Generate tailored resume and cover letter .docx files")
    parser.add_argument("--config", required=True,
                        help="Path to config.json with personal info and paths")
    args = parser.parse_args()

    cfg = load_config(args.config)
    output_dir = Path(cfg["output_dir"])

    resume_json = output_dir / "resume_tailored.json"
    cover_json = output_dir / "cover_letter_tailored.json"

    if not resume_json.exists():
        print(f"ERROR: {resume_json} not found", file=sys.stderr)
        sys.exit(1)
    if not cover_json.exists():
        print(f"ERROR: {cover_json} not found", file=sys.stderr)
        sys.exit(1)

    with open(resume_json, encoding="utf-8") as f:
        resume_data = json.load(f)
    with open(cover_json, encoding="utf-8") as f:
        cover_data = json.load(f)

    company = resume_data["company"]
    short_title = resume_data["short_title"]

    # Build filename-safe name (e.g. "Tom Colarusso" -> "Tom_Colarusso")
    safe_name = cfg["name"].replace(" ", "_")

    resume_doc = build_resume(resume_data, cfg)
    resume_path = output_dir / f"{safe_name}_Resume_{company}_{short_title}.docx"
    resume_doc.save(str(resume_path))
    print(f"Resume saved: {resume_path}")

    cover_doc = build_cover_letter(cover_data, cfg)
    cover_path = output_dir / f"{safe_name}_CoverLetter_{company}_{short_title}.docx"
    cover_doc.save(str(cover_path))
    print(f"Cover letter saved: {cover_path}")


if __name__ == "__main__":
    main()
